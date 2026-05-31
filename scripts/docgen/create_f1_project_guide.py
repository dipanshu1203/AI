from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "projects" / "F1" / "docs"
VISUALS = DOCS / "visuals"
OUTPUT = DOCS / "F1_Project_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LINE = "DADCE0"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def font(size=24, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill, outline="#C7CDD4", text_fill="#111111", width=2):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=width)
    x1, y1, x2, y2 = xy
    max_width = x2 - x1 - 28
    words = text.split()
    lines = []
    current = ""
    f = font(24, bold=True)
    for word in words:
        probe = f"{current} {word}".strip()
        if draw.textlength(probe, font=f) <= max_width:
            current = probe
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    total_h = len(lines) * 29
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        draw.text((x1 + 14, y), line, fill=text_fill, font=f)
        y += 29


def arrow(draw, start, end, color="#5F6368"):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        points = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)] if ex > sx else [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    else:
        points = [(ex, ey), (ex - 8, ey - 14), (ex + 8, ey - 14)] if ey > sy else [(ex, ey), (ex - 8, ey + 14), (ex + 8, ey + 14)]
    draw.polygon(points, fill=color)


def make_deploy_visual(path):
    img = Image.new("RGB", (1500, 780), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = "Public URL Path"
    draw.text((60, 46), title, fill="#111111", font=font(40, bold=True))
    draw.text((60, 96), "The GitHub repo name supplies /AI/. The project folder supplies /projects/F1/.", fill="#555555", font=font(24))

    boxes = [
        ((80, 210, 360, 330), "Local folder projects/F1", "#F8FAFC"),
        ((470, 210, 750, 330), "Push to GitHub repo AI", "#F8FAFC"),
        ((860, 210, 1140, 330), "GitHub Pages builds main", "#F8FAFC"),
        ((430, 480, 1070, 610), "https://dipanshu1203.github.io/AI/projects/F1/", "#EAF4FF"),
    ]
    for xy, text, fill in boxes:
        rounded_box(draw, xy, text, fill)
    arrow(draw, (360, 270), (470, 270))
    arrow(draw, (750, 270), (860, 270))
    arrow(draw, (1000, 330), (870, 480))

    img.save(path)


def make_refresh_visual(path):
    img = Image.new("RGB", (1500, 840), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 46), "Prediction Refresh Flow", fill="#111111", font=font(40, bold=True))
    draw.text((60, 96), "A scheduled GitHub Action updates prediction.json, then Pages publishes the refreshed data.", fill="#555555", font=font(24))

    boxes = [
        ((80, 210, 360, 330), "Monday schedule or manual run", "#F8FAFC"),
        ((470, 210, 750, 330), "Node updater calls Jolpica-F1", "#F8FAFC"),
        ((860, 210, 1140, 330), "prediction.json is committed", "#F8FAFC"),
        ((470, 500, 750, 620), "GitHub Pages rebuilds", "#F8FAFC"),
        ((860, 500, 1140, 620), "Website fetches latest JSON", "#EAF4FF"),
    ]
    for xy, text, fill in boxes:
        rounded_box(draw, xy, text, fill)
    arrow(draw, (360, 270), (470, 270))
    arrow(draw, (750, 270), (860, 270))
    arrow(draw, (1000, 330), (650, 500))
    arrow(draw, (750, 560), (860, 560))

    img.save(path)


def set_cell(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if bold:
        run.font.color.rgb = RGBColor(17, 17, 17)
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                elem = OxmlElement(f"w:{side}")
                elem.set(qn("w:w"), "120" if side in ("start", "end") else "80")
                elem.set(qn("w:type"), "dxa")
                tc_mar.append(elem)
            tc_pr.append(tc_mar)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell(cell, text, fill=CALLOUT_FILL)
    doc.add_paragraph()


def build_doc():
    VISUALS.mkdir(parents=True, exist_ok=True)
    deploy_img = VISUALS / "deploy-path.png"
    refresh_img = VISUALS / "prediction-refresh-flow.png"
    make_deploy_visual(deploy_img)
    make_refresh_visual(refresh_img)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("F1 Website Project Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 17, 17)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Build, deploy, and automated prediction refresh steps").font.color.rgb = MUTED
    doc.add_paragraph("Repository path: projects/F1")
    doc.add_paragraph("Public URL: https://dipanshu1203.github.io/AI/projects/F1/")

    add_callout(
        doc,
        "Purpose: This guide documents how the simple F1 site was built, where the files live, how it is deployed on GitHub Pages, and how the prediction data updates after race weekends."
    )

    add_heading(doc, "Project Snapshot", 1)
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    set_cell(table.cell(0, 0), "Area", bold=True, fill=HEADER_FILL)
    set_cell(table.cell(0, 1), "Details", bold=True, fill=HEADER_FILL)
    rows = [
        ("Site type", "Static HTML, CSS, and JavaScript"),
        ("Main folder", "projects/F1"),
        ("Homepage", "projects/F1/index.html"),
        ("Prediction data", "projects/F1/data/prediction.json"),
        ("Automation", ".github/workflows/update-f1-prediction.yml"),
        ("Updater script", "scripts/update-f1-prediction.mjs"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell(cells[0], left, bold=True)
        set_cell(cells[1], right)

    add_heading(doc, "What We Built", 1)
    add_bullet(doc, "A minimal Formula 1 dashboard with upcoming Grand Prix calendar, standings, news links, and prediction widget.")
    add_bullet(doc, "Driver portraits and team car images are pulled from Formula 1 media URLs.")
    add_bullet(doc, "The prediction model weighs current team form and last year track-result history.")
    add_bullet(doc, "GitHub Pages publishes the site publicly at the requested URL.")
    add_bullet(doc, "A scheduled GitHub Action refreshes prediction data from the Jolpica-F1 API.")

    add_heading(doc, "Folder Structure", 1)
    structure = doc.add_paragraph()
    structure.style = "No Spacing"
    structure.add_run(
        "projects/F1/\n"
        "  index.html\n"
        "  styles.css\n"
        "  script.js\n"
        "  assets/f1-hero.png\n"
        "  data/prediction.json\n"
        "  docs/F1_Project_Guide.docx\n"
        "  docs/visuals/\n"
        ".github/workflows/update-f1-prediction.yml\n"
        "scripts/update-f1-prediction.mjs"
    ).font.name = "Courier New"

    add_heading(doc, "Deployment Path", 1)
    doc.add_picture(str(deploy_img), width=Inches(6.4))
    p = doc.add_paragraph()
    p.add_run("Why this URL works: ").bold = True
    p.add_run("GitHub Pages adds the repository name, AI, to the base URL. The project folder is projects/F1, so the final path becomes /AI/projects/F1/.")

    add_heading(doc, "Build and Publish Steps", 1)
    steps = [
        "Create the static site in projects/F1 with index.html, styles.css, script.js, and local assets.",
        "Verify locally at http://localhost:8000/projects/F1/ using a simple static server.",
        "Commit the site files to the main branch.",
        "Push main to https://github.com/dipanshu1203/AI.git.",
        "Enable GitHub Pages from the main branch root.",
        "Open https://dipanshu1203.github.io/AI/projects/F1/ and confirm it returns HTTP 200.",
    ]
    for step in steps:
        add_number(doc, step)

    add_heading(doc, "Prediction Refresh Flow", 1)
    doc.add_picture(str(refresh_img), width=Inches(6.4))
    add_bullet(doc, "The workflow runs every Monday at 08:30 UTC and can also be started manually from GitHub Actions.")
    add_bullet(doc, "The Node script calls Jolpica-F1 for schedule, current constructor standings, and last season result at the same circuit.")
    add_bullet(doc, "The script rewrites projects/F1/data/prediction.json when the model changes.")
    add_bullet(doc, "The GitHub Action commits the updated JSON and pushes it back to main.")
    add_bullet(doc, "GitHub Pages rebuilds, and the browser fetches the latest JSON on page load.")

    add_heading(doc, "Manual Commands", 1)
    commands = [
        ("Run locally", "python3 -m http.server 8000"),
        ("Open local site", "http://localhost:8000/projects/F1/"),
        ("Refresh prediction data locally", "node scripts/update-f1-prediction.mjs"),
        ("Push updates", "git add . && git commit -m \"Update F1 project\" && git push origin main"),
        ("Trigger workflow manually", "gh workflow run update-f1-prediction.yml --repo dipanshu1203/AI"),
    ]
    cmd_table = doc.add_table(rows=1, cols=2)
    style_table(cmd_table)
    set_cell(cmd_table.cell(0, 0), "Task", bold=True, fill=HEADER_FILL)
    set_cell(cmd_table.cell(0, 1), "Command or URL", bold=True, fill=HEADER_FILL)
    for task, command in commands:
        cells = cmd_table.add_row().cells
        set_cell(cells[0], task, bold=True)
        set_cell(cells[1], command)

    add_heading(doc, "How To Maintain It", 1)
    add_bullet(doc, "Edit page layout/content in index.html and styles.css.")
    add_bullet(doc, "Edit prediction UI behavior in script.js.")
    add_bullet(doc, "Change model scoring logic in scripts/update-f1-prediction.mjs.")
    add_bullet(doc, "Keep prediction.json committed because the static site reads it directly from GitHub Pages.")
    add_bullet(doc, "If the public URL changes, check the repository name and project folder path first.")

    add_heading(doc, "Verification Checklist", 1)
    check_table = doc.add_table(rows=1, cols=3)
    style_table(check_table)
    for idx, label in enumerate(["Check", "How", "Expected result"]):
        set_cell(check_table.cell(0, idx), label, bold=True, fill=HEADER_FILL)
    checks = [
        ("Local page", "Open /projects/F1/", "Hero, standings, prediction, and news render"),
        ("Images", "Inspect page or browser console", "No broken driver/team images"),
        ("Prediction JSON", "Open /projects/F1/data/prediction.json", "HTTP 200 and valid JSON"),
        ("GitHub Action", "GitHub Actions run history", "Update F1 Prediction Data succeeds"),
        ("Public page", "Open public URL", "HTTP 200 and current site visible"),
    ]
    for row in checks:
        cells = check_table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value, bold=(idx == 0))

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("F1 Website Project Guide").font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
